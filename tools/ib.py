from docx import Document


def convert_to_html_tags(run):
    text = run.text
    if run.italic:
        text = '<i>' + text + '</i>'
    if run.bold:
        text = '<b>' + text + '</b>'
    return text.replace('\n', '<br/>').replace('\r', '<br/>')


def convert_word_to_html(file_path):
    document = Document(file_path)
    paragraphs = document.paragraphs

    html_content = ''
    for paragraph in paragraphs:
        for run in paragraph.runs:
            converted_text = convert_to_html_tags(run)
            html_content += converted_text

    return html_content


def load(target_file, content):
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(content)
    document.save(target_file)


# Exemple d'utilisation
file_path = 'tools/scipt.docx'
html_content = convert_word_to_html(file_path)
load('tools/bon.docx', html_content)
